import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog, TclError
import re
import nltk
from collections import Counter
import numpy as np
import random
import os
import threading
from queue import Queue
import hashlib
import docx
import pdfplumber
from transformers import pipeline, AutoTokenizer
import torch
import torch.nn as nn
from huggingface_hub import login
import logging
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer, util
from textblob import TextBlob
from tkinterdnd2 import TkinterDnD, DND_FILES

# Suppress oneDNN warnings
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class NLTKManager:
    """Handles NLTK resource initialization"""
    REQUIRED_PACKAGES = ['punkt', 'stopwords', 'averaged_perceptron_tagger', 'wordnet']
    
    @classmethod
    def initialize(cls):
        """Initialize all required NLTK resources"""
        nltk_data_path = os.path.join(os.path.expanduser("~"), "nltk_data")
        os.makedirs(nltk_data_path, exist_ok=True)
        nltk.data.path.append(nltk_data_path)
        
        for package in cls.REQUIRED_PACKAGES:
            try:
                nltk.data.find(package)
            except LookupError:
                try:
                    nltk.download(package, quiet=True)
                except Exception as e:
                    logging.warning(f"Failed to download NLTK package {package}: {e}")

class AIModelPatterns:
    """Contains patterns specific to different AI models"""
    PATTERNS = {
        'gpt4': {
            'phrases': [
                'as an ai language model', 'as a large language model',
                'i don\'t have personal opinions', 'my training data only goes up to',
                'i strive to be helpful', 'in accordance with my ethical guidelines'
            ],
            'sentence_structure': {
                'avg_length': range(18, 26),
                'transition_freq': 0.15,
                'intro_clause_prob': 0.7
            }
        },
        'gpt3': {
            'phrases': [
                'as an ai', 'i don\'t have consciousness',
                'my knowledge is limited to', 'i was trained on'
            ],
            'sentence_structure': {
                'avg_length': range(15, 21),
                'transition_freq': 0.2,
                'intro_clause_prob': 0.6
            }
        },
        'claude': {
            'phrases': [
                'as an ai assistant', 'i aim to be helpful',
                'while i don\'t have subjective experiences',
                'my responses are based on patterns in my training data'
            ],
            'sentence_structure': {
                'avg_length': range(20, 29),
                'transition_freq': 0.12,
                'intro_clause_prob': 0.65
            }
        },
        'deepseek': {
            'phrases': [
                'based on my analysis', 'the data suggests',
                'statistically speaking', 'from an objective standpoint'
            ],
            'sentence_structure': {
                'avg_length': range(16, 23),
                'transition_freq': 0.18,
                'intro_clause_prob': 0.55
            }
        }
    }

    COMMON_AI_PATTERNS = [
        r'\b(however|therefore|moreover|furthermore)\b',
        r'\b(it is (important|worth noting|crucial) that)\b',
        r'\b(in (conclusion|summary|addition))\b',
        r'\b(as (mentioned|stated) (previously|above))\b',
        r'\b(this (suggests|indicates|implies))\b'
    ]

class ANN(nn.Module):
    """Simple ANN for unsupervised learning of text embeddings"""
    def __init__(self, input_dim=384, hidden_dim=256):
        super(ANN, self).__init__()
        self.encoder = nn.Sequential(
            nn.Linear(input_dim, hidden_dim),
            nn.ReLU(),
            nn.Linear(hidden_dim, hidden_dim),
            nn.ReLU(),
            nn.Linear(hidden_dim, input_dim)
        )
    
    def forward(self, x):
        return self.encoder(x)

class AdvancedParaphraser:
    """Manages multiple paraphrasing models and ANN for advanced paraphrasing"""
    def __init__(self, device):
        self.device = device
        self.models = {}
        self.tokenizers = {}
        self.sentence_encoder = SentenceTransformer('all-MiniLM-L6-v2', device=device)
        self.ann = ANN(input_dim=384, hidden_dim=256).to(device)
        self.optimizer = torch.optim.Adam(self.ann.parameters(), lr=0.001)
        
        # Load transformer models
        model_configs = [
            ('parrot', 'prithivida/parrot_paraphraser_on_T5', 'text2text-generation'),
            ('t5', 't5-base', 'text2text-generation'),
            ('bart', 'facebook/bart-large-cnn', 'summarization')
        ]
        
        for name, model_id, task in model_configs:
            try:
                logging.info(f"Loading {name} model: {model_id}")
                self.models[name] = pipeline(task, model=model_id, device=device)
                self.tokenizers[name] = AutoTokenizer.from_pretrained(model_id)
                logging.info(f"{name} model loaded successfully")
            except Exception as e:
                logging.error(f"Failed to load {name} model: {e}")
                self.models[name] = None
    
    def paraphrase(self, text, diversity=0.7):
        """Paraphrase text using ensemble of models and ANN"""
        if not text.strip():
            return "Error: No text provided"
        
        chunks = self._chunk_text(text)
        paraphrased_chunks = []
        
        for chunk in chunks:
            candidates = []
            for name, model in self.models.items():
                if model is None:
                    continue
                try:
                    if name == 'parrot':
                        result = model(f"paraphrase: {chunk}", max_length=512, num_return_sequences=1)
                        candidates.append(result[0]['generated_text'])
                    elif name == 't5':
                        result = model(f"paraphrase: {chunk}", max_length=512)
                        candidates.append(result[0]['generated_text'])
                    elif name == 'bart':
                        result = model(chunk, max_length=150, min_length=30, do_sample=True)
                        candidates.append(result[0]['summary_text'])
                except Exception as e:
                    logging.error(f"Error in {name} paraphrasing: {e}")
            
            # Enhance with ANN
            best_candidate = self._select_best_paraphrase(chunk, candidates, diversity)
            if best_candidate:
                embedding = self.sentence_encoder.encode(best_candidate, convert_to_tensor=True)
                refined_embedding = self.ann(embedding)
                paraphrased_chunks.append(best_candidate)
            else:
                paraphrased_chunks.append(chunk)
        
        return ' '.join(paraphrased_chunks)
    
    def _chunk_text(self, text, chunk_size=512):
        """Split text into chunks for processing"""
        try:
            tokens = nltk.word_tokenize(text)
            return [' '.join(tokens[i:i+chunk_size]) for i in range(0, len(tokens), chunk_size)]
        except Exception as e:
            logging.error(f"Error chunking text: {e}")
            return [text]
    
    def _select_best_paraphrase(self, original, candidates, diversity=0.7):
        """Select the best paraphrase based on diversity and similarity"""
        if not candidates:
            return original
        
        original_embedding = self.sentence_encoder.encode(original, convert_to_tensor=True)
        candidate_embeddings = self.sentence_encoder.encode(candidates, convert_to_tensor=True)
        
        similarities = util.cos_sim(original_embedding, candidate_embeddings).flatten()
        target_similarity = 1.0 - diversity
        valid_indices = [i for i, sim in enumerate(similarities) if target_similarity - 0.1 <= sim <= target_similarity + 0.1]
        
        if not valid_indices:
            return candidates[0] if candidates else original
        return candidates[max(valid_indices, key=lambda i: similarities[i])]
    
    def train_on_documents(self, texts):
        """Train ANN on human-written document embeddings"""
        if not texts:
            return
        
        self.ann.train()
        embeddings = self.sentence_encoder.encode(texts, convert_to_tensor=True, show_progress_bar=False)
        for epoch in range(5):  # Limited epochs for quick training
            self.optimizer.zero_grad()
            output = self.ann(embeddings)
            loss = nn.MSELoss()(output, embeddings)  # Autoencoder-style training
            loss.backward()
            self.optimizer.step()
            logging.info(f"ANN training epoch {epoch+1}, loss: {loss.item()}")
        self.ann.eval()
    
    def save(self, path):
        """Save the ensemble pipeline and ANN"""
        os.makedirs(path, exist_ok=True)
        for name, model in self.models.items():
            if model is not None:
                try:
                    model.save_pretrained(os.path.join(path, name))
                    self.tokenizers[name].save_pretrained(os.path.join(path, name))
                    logging.info(f"Saved {name} model to {path}")
                except Exception as e:
                    logging.error(f"Failed to save {name} model: {e}")
        try:
            torch.save(self.ann.state_dict(), os.path.join(path, "ann.pt"))
            logging.info(f"Saved ANN to {path}")
        except Exception as e:
            logging.error(f"Failed to save ANN: {e}")
    
    def load(self, path):
        """Load the ensemble pipeline and ANN"""
        for name, model in self.models.items():
            try:
                model_path = os.path.join(path, name)
                if os.path.exists(model_path):
                    task = 'text2text-generation' if name in ['parrot', 't5'] else 'summarization'
                    self.models[name] = pipeline(task, model=model_path, device=self.device)
                    self.tokenizers[name] = AutoTokenizer.from_pretrained(model_path)
                    logging.info(f"Loaded {name} model from {path}")
                else:
                    logging.warning(f"Model path {model_path} does not exist")
            except Exception as e:
                logging.error(f"Failed to load {name} model: {e}")
                self.models[name] = None
        try:
            ann_path = os.path.join(path, "ann.pt")
            if os.path.exists(ann_path):
                self.ann.load_state_dict(torch.load(ann_path, map_location=self.device))
                logging.info(f"Loaded ANN from {path}")
        except Exception as e:
            logging.error(f"Failed to load ANN: {e}")

class AIContentDetector:
    """Advanced AI content detector with model-specific pattern recognition and pretrained models"""
    def __init__(self):
        # Authenticate with Hugging Face
        try:
            login(token="HUGGING FACE TOKEN")
            logging.info("Hugging Face login successful")
        except Exception as e:
            logging.error(f"Hugging Face login failed: {e}")
        
        self._init_patterns()
        self._init_humanization()
        self.cache = {}
        NLTKManager.initialize()
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        logging.info(f"Using device: {self.device}")
        
        # Initialize pretrained models
        self.ai_detector = None
        self.perplexity_model = None
        self.paraphraser = None
        try:
            logging.info("Loading pretrained models...")
            self.ai_detector = pipeline("text-classification", model="roberta-base-openai-detector", device=self.device)
            self.perplexity_model = pipeline("text-generation", model="gpt2", device=self.device)
            self.paraphraser = AdvancedParaphraser(self.device)
            logging.info("Pretrained models loaded successfully")
        except Exception as e:
            logging.error(f"Error loading models: {e}. Falling back to rule-based detection.")
            messagebox.showwarning("Warning", "Failed to load pretrained models. Using rule-based detection.")
        
        # Load saved models if available
        self.load_models("./advanced_model")
        # Save models for future use
        self.save_models("./advanced_model")

    def _init_patterns(self):
        """Initialize detection patterns and compile regex"""
        self.model_patterns = AIModelPatterns.PATTERNS
        self.common_ai_regex = [re.compile(pattern, re.IGNORECASE) for pattern in AIModelPatterns.COMMON_AI_PATTERNS]
        self.passive_voice = re.compile(r'\b(am|is|are|was|were|be|being|been)\s+[a-z]+ed\b', re.IGNORECASE)
        self.nominalizations = re.compile(r'\b\w+(tion|ment|ance|ence|ity)\b', re.IGNORECASE)
        self.hedging = re.compile(r'\b(may|might|could|would|should|possibly|perhaps)\b', re.IGNORECASE)

    def _init_humanization(self):
        """Initialize humanization templates"""
        self.templates = {
            'as an ai': ['from my perspective', 'in my view', 'I believe'],
            'however': ['but', 'though', 'on the other hand'],
            'therefore': ['so', 'thus', 'that\'s why'],
            'in conclusion': ['to wrap up', 'finally', 'all things considered']
        }
        self.contractions = {
            'do not': 'don\'t', 'does not': 'doesn\'t', 'did not': 'didn\'t',
            'will not': 'won\'t', 'cannot': 'can\'t', 'are not': 'aren\'t'
        }
        self.interjections = ['well', 'you know', 'actually', 'I mean', 'honestly']
        self.synonyms = {
            'good': ['great', 'excellent', 'fantastic', 'superb'],
            'bad': ['poor', 'terrible', 'awful', 'subpar'],
            'important': ['crucial', 'vital', 'essential', 'key']
        }

    def save_models(self, path):
        """Save all models"""
        if self.ai_detector:
            try:
                self.ai_detector.save_pretrained(os.path.join(path, "ai_detector"))
                logging.info(f"Saved AI detector to {path}")
            except Exception as e:
                logging.error(f"Failed to save AI detector: {e}")
        if self.perplexity_model:
            try:
                self.perplexity_model.save_pretrained(os.path.join(path, "perplexity_model"))
                logging.info(f"Saved perplexity model to {path}")
            except Exception as e:
                logging.error(f"Failed to save perplexity model: {e}")
        if self.paraphraser:
            self.paraphraser.save(path)

    def load_models(self, path):
        """Load all models"""
        try:
            detector_path = os.path.join(path, "ai_detector")
            if os.path.exists(detector_path):
                self.ai_detector = pipeline("text-classification", model=detector_path, device=self.device)
                logging.info(f"Loaded AI detector from {path}")
            else:
                logging.warning(f"AI detector path {detector_path} does not exist")
        except Exception as e:
            logging.error(f"Failed to load AI detector: {e}")
        try:
            perplexity_path = os.path.join(path, "perplexity_model")
            if os.path.exists(perplexity_path):
                self.perplexity_model = pipeline("text-generation", model=perplexity_path, device=self.device)
                logging.info(f"Loaded perplexity model from {path}")
            else:
                logging.warning(f"Perplexity model path {perplexity_path} does not exist")
        except Exception as e:
            logging.error(f"Failed to load perplexity model: {e}")
        if self.paraphraser:
            self.paraphraser.load(path)

    def learn_from_documents(self, texts):
        """Train the paraphraser on human-written documents"""
        if not texts:
            return
        try:
            self.paraphraser.train_on_documents(texts)
            logging.info("ANN trained on uploaded documents")
        except Exception as e:
            logging.error(f"Error training on documents: {e}")

    def _chunk_text(self, text, chunk_size=512):
        """Split text into chunks for processing"""
        try:
            tokens = nltk.word_tokenize(text)
            return [' '.join(tokens[i:i+chunk_size]) for i in range(0, len(tokens), chunk_size)]
        except Exception as e:
            logging.error(f"Error chunking text: {e}")
            return [text]

    def _calculate_perplexity(self, text):
        """Calculate perplexity using a language model"""
        if not self.perplexity_model:
            return 0.0
        try:
            encodings = self.perplexity_model.tokenizer(text, return_tensors='pt').to(self.device)
            input_ids = encodings['input_ids']
            with torch.no_grad():
                outputs = self.perplexity_model.model(input_ids, labels=input_ids)
            return float(torch.exp(outputs.loss).item())
        except Exception as e:
            logging.error(f"Error calculating perplexity: {e}")
            return 0.0

    def _ngram_repetition(self, text, n=3):
        """Calculate n-gram repetition score"""
        words = nltk.word_tokenize(text.lower())
        if len(words) < n:
            return 0.0
        ngrams = [tuple(words[i:i+n]) for i in range(len(words)-n+1)]
        ngram_counts = Counter(ngrams)
        repeated = sum(count > 1 for count in ngram_counts.values())
        return min(1.0, repeated / max(1, len(ngrams)))

    def _stylometric_features(self, text):
        """Calculate stylometric features"""
        blob = TextBlob(text)
        sentences = blob.sentences
        if not sentences:
            return 0.0
        complexity = sum(len(sent.words) for sent in sentences) / len(sentences)
        punctuation = len(re.findall(r'[.,!?]', text)) / len(text)
        return min(1.0, (complexity / 20 + punctuation) / 2)

    def detect_ai_model(self, text):
        """Detect AI-generated text using pretrained model and advanced algorithms"""
        if not text.strip():
            logging.warning("Empty text provided for AI detection")
            return {'error': 'Empty text provided'}
        
        scores = {model: 0.0 for model in self.model_patterns}
        text_lower = text.lower()
        
        # Rule-based pattern detection
        for model, patterns in self.model_patterns.items():
            for phrase in patterns['phrases']:
                if phrase in text_lower:
                    scores[model] += 2.0
            sentences = nltk.sent_tokenize(text)
            if sentences:
                avg_len = sum(len(nltk.word_tokenize(s)) for s in sentences) / len(sentences)
                model_avg = (min(self.model_patterns[model]['sentence_structure']['avg_length']) +
                            max(self.model_patterns[model]['sentence_structure']['avg_length'])) / 2
                scores[model] += 1.0 - min(1.0, abs(avg_len - model_avg) / 10)
        
        for pattern in self.common_ai_regex:
            if pattern.search(text):
                for model in scores:
                    scores[model] += 0.2
        
        # Advanced detection algorithms
        perplexity_score = self._calculate_perplexity(text)
        ngram_score = self._ngram_repetition(text)
        style_score = self._stylometric_features(text)
        
        # Pretrained model detection
        avg_ai_score = 0.0
        if self.ai_detector:
            chunks = self._chunk_text(text, chunk_size=512)
            ai_scores = []
            for chunk in chunks:
                try:
                    result = self.ai_detector(chunk, truncation=True, max_length=512)
                    score = result[0]['score'] if result[0]['label'].startswith('AI') else 1.0 - result[0]['score']
                    ai_scores.append(score)
                except Exception as e:
                    logging.error(f"Error in AI detection for chunk: {e}")
            avg_ai_score = np.mean(ai_scores) * 100 if ai_scores else 0.0
        else:
            logging.warning("AI detector not loaded, using rule-based scores only")
            avg_ai_score = sum(scores.values()) / len(scores) if scores else 0.0

        # Combine scores
        max_score = max(scores.values()) if scores else 1.0
        if max_score > 0:
            for model in scores:
                scores[model] = min(100.0, (scores[model] * 0.3 + avg_ai_score * 0.3 + 
                                           perplexity_score * 0.2 + ngram_score * 0.1 + style_score * 0.1) * 100 / max_score)
        
        return scores

    def analyze_text(self, text):
        """Comprehensive AI content analysis"""
        if not text.strip():
            return {'error': 'Empty text provided'}
        
        text_hash = hashlib.md5(text.encode()).hexdigest()
        if text_hash in self.cache:
            logging.info("Returning cached analysis")
            return self.cache[text_hash]
        
        model_scores = self.detect_ai_model(text)
        if 'error' in model_scores:
            return model_scores
        
        likely_model = max(model_scores, key=model_scores.get) if model_scores else None
        
        analysis = {
            'model_scores': model_scores,
            'likely_model': likely_model,
            'scores': {
                'passive_voice': self._score_passive_voice(text),
                'nominalizations': self._score_nominalizations(text),
                'hedging': self._score_hedging(text),
                'sentence_uniformity': self._score_sentence_uniformity(text),
                'lexical_diversity': self._score_lexical_diversity(text),
                'common_ai_patterns': self._score_common_patterns(text),
                'perplexity': self._calculate_perplexity(text),
                'ngram_repetition': self._ngram_repetition(text),
                'stylometric': self._stylometric_features(text)
            }
        }
        
        weights = {
            'passive_voice': 0.1,
            'nominalizations': 0.1,
            'hedging': 0.1,
            'sentence_uniformity': 0.2,
            'lexical_diversity': 0.15,
            'common_ai_patterns': 0.15,
            'perplexity': 0.1,
            'ngram_repetition': 0.05,
            'stylometric': 0.05
        }
        
        weighted_sum = sum(analysis['scores'][k] * weights[k] for k in weights)
        analysis['ai_probability'] = min(100.0, weighted_sum * 100)
        analysis['suggestions'] = self._generate_suggestions(analysis)
        
        self.cache[text_hash] = analysis
        return analysis

    def _score_passive_voice(self, text):
        """Score passive voice usage"""
        matches = len(self.passive_voice.findall(text.lower()))
        words = len(nltk.word_tokenize(text))
        return min(1.0, matches / max(1, words/20))

    def _score_nominalizations(self, text):
        """Score nominalization usage"""
        matches = len(self.nominalizations.findall(text.lower()))
        words = len(nltk.word_tokenize(text))
        return min(1.0, matches / max(1, words/15))

    def _score_hedging(self, text):
        """Score hedging language usage"""
        matches = len(self.hedging.findall(text.lower()))
        words = len(nltk.word_tokenize(text))
        return min(1.0, matches / max(1, words/25))

    def _score_sentence_uniformity(self, text):
        """Score sentence structure uniformity"""
        sentences = nltk.sent_tokenize(text)
        if len(sentences) < 3:
            return 0.0
        lengths = [len(nltk.word_tokenize(s)) for s in sentences]
        avg_len = np.mean(lengths)
        std_dev = np.std(lengths)
        return max(0.0, 1.0 - (std_dev / avg_len if avg_len > 0 else 0.0))

    def _score_lexical_diversity(self, text):
        """Score lexical diversity"""
        words = [w.lower() for w in nltk.word_tokenize(text) if w.isalpha()]
        if not words:
            return 0.0
        unique = len(set(words))
        return max(0.0, 1.0 - (unique / len(words)))

    def _score_common_patterns(self, text):
        """Score common AI patterns"""
        matches = sum(1 for pattern in self.common_ai_regex if pattern.search(text))
        sentences = len(nltk.sent_tokenize(text))
        return min(1.0, matches / max(1, sentences/3))

    def _generate_suggestions(self, analysis):
        """Generate targeted suggestions based on analysis"""
        suggestions = []
        scores = analysis['scores']
        
        if analysis['ai_probability'] > 70:
            suggestions.append(f"High AI probability detected (likely {analysis['likely_model']})")
        if scores['passive_voice'] > 0.3:
            suggestions.append("Reduce passive voice constructions")
        if scores['nominalizations'] > 0.25:
            suggestions.append("Convert nominalizations to active verbs")
        if scores['hedging'] > 0.2:
            suggestions.append("Reduce hedging language (may, might, could)")
        if scores['sentence_uniformity'] > 0.7:
            suggestions.append("Vary sentence lengths and structures")
        if scores['lexical_diversity'] > 0.6:
            suggestions.append("Increase vocabulary diversity")
        if scores['common_ai_patterns'] > 0.5:
            suggestions.append("Avoid common AI phrasing patterns")
        if scores['perplexity'] > 100:
            suggestions.append("Simplify language to reduce high perplexity")
        if scores['ngram_repetition'] > 0.3:
            suggestions.append("Reduce repetitive n-grams")
        if scores['stylometric'] > 0.5:
            suggestions.append("Adjust stylistic features for naturalness")
        
        return suggestions[:8]

    def humanize_text(self, text, diversity=0.7):
        """Convert AI-generated text to human-like writing"""
        if not text.strip():
            return text
        
        # Paraphrase using the advanced paraphraser
        paraphrased = self.paraphrase_text(text, diversity)
        if paraphrased.startswith("Error:"):
            paraphrased = text
        
        # Apply rule-based humanization
        model_scores = self.detect_ai_model(paraphrased)
        likely_model = max(model_scores, key=model_scores.get) if model_scores and 'error' not in model_scores else None
        
        if likely_model:
            for phrase in self.model_patterns[likely_model]['phrases']:
                if phrase.lower() in paraphrased.lower():
                    replacement = random.choice(self.templates.get(phrase.lower(), ['']))
                    paraphrased = re.sub(phrase, replacement, paraphrased, flags=re.IGNORECASE)
        
        for pattern, replacements in self.templates.items():
            if pattern.lower() in paraphrased.lower():
                paraphrased = re.sub(
                    r'\b' + pattern + r'\b',
                    lambda m: random.choice(replacements),
                    paraphrased,
                    flags=re.IGNORECASE
                )
        
        for full, contracted in self.contractions.items():
            paraphrased = re.sub(r'\b' + full + r'\b', contracted, paraphrased, flags=re.IGNORECASE)
        
        for word, synonyms in self.synonyms.items():
            paraphrased = re.sub(r'\b' + word + r'\b', lambda m: random.choice(synonyms), paraphrased, flags=re.IGNORECASE)
        
        sentences = nltk.sent_tokenize(paraphrased)
        for i in range(len(sentences)):
            if random.random() < 0.4 and i > 0:
                sentences[i] = random.choice(self.interjections).capitalize() + ', ' + sentences[i].lower()
            if random.random() < 0.3 and not sentences[i].endswith('?'):
                sentences[i] = sentences[i].rstrip('.') + random.choice([' right?', ' you know?', ' yeah?'])
            if random.random() < 0.2:
                blob = TextBlob(sentences[i])
                sentences[i] = str(blob.correct())  # Basic spelling/grammar correction
        
        humanized = ' '.join(sentences)
        humanized = re.sub(r'([.!?])\s+([a-z])', lambda m: m.group(1) + ' ' + m.group(2).upper(), humanized)
        if humanized:
            humanized = humanized[0].upper() + humanized[1:]
        
        return humanized

    def summarize_text(self, text):
        """Summarize text using pretrained model"""
        if not text.strip():
            return "Error: No text provided"
        if not self.paraphraser or not self.paraphraser.models.get('bart'):
            return "Error: Summarizer not available, please check model loading"
        
        chunks = self._chunk_text(text, chunk_size=512)
        summaries = []
        for chunk in chunks:
            try:
                summary = self.paraphraser.models['bart'](chunk, max_length=150, min_length=30, do_sample=False)
                summaries.append(summary[0]['summary_text'])
            except Exception as e:
                logging.error(f"Error in summarization: {e}")
                summaries.append("Error: Failed to summarize chunk")
        return ' '.join(summaries)

    def paraphrase_text(self, text, diversity=0.7):
        """Paraphrase text using advanced paraphraser"""
        if not text.strip():
            return "Error: No text provided"
        if not self.paraphraser:
            return "Error: Paraphraser not available, please check model loading"
        
        return self.paraphraser.paraphrase(text, diversity)

class AsyncTaskManager:
    """Handles background tasks for responsive UI"""
    def __init__(self, log_callback):
        self.queue = Queue()
        self.running = True
        self.log_callback = log_callback
        self.thread = threading.Thread(target=self._worker, daemon=True)
        self.thread.start()
    
    def _worker(self):
        while self.running:
            task, callback = self.queue.get()
            try:
                result = task()
                if callback:
                    callback(result)
            except Exception as e:
                logging.error(f"Task failed: {e}")
                self.log_callback(f"Task failed: {str(e)}")
            self.queue.task_done()
    
    def add_task(self, task, callback=None):
        self.queue.put((task, callback))
    
    def shutdown(self):
        self.running = False
        self.thread.join()

class AIDetectorGUI:
    """Advanced GUI for AI content detection and processing with modern white theme"""
    def __init__(self):
        self.root = TkinterDnD.Tk()
        self.root.title("Advanced AI Content Detector")
        self.root.geometry("1280x720")
        self.detector = AIContentDetector()
        self.task_manager = AsyncTaskManager(self._log_message)
        self.file_queue = []
        self.history = []
        self.progress_animation_id = None
        self._setup_ui()
        
    def _setup_ui(self):
        """Setup the enhanced GUI with white theme and improved layout"""
        # Style configuration
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TFrame', background='#ffffff')
        style.configure('TLabel', background='#ffffff', foreground='#333333')
        style.configure('TButton', background='#007bff', foreground='#ffffff', padding=8)
        style.configure('TNotebook', background='#f8f9fa')
        style.configure('TNotebook.Tab', background='#e9ecef', foreground='#333333', padding=(10, 5))
        style.map('TButton', background=[('active', '#0056b3')])
        style.configure('TProgressbar', background='#28a745', troughcolor='#dee2e6')
        
        # Main container
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Notebook for tabs
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Input Tab
        input_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(input_frame, text="Input")
        
        input_label = ttk.Label(input_frame, text="Enter or Drop Text/Documents", font=('Helvetica', 12, 'bold'))
        input_label.pack(anchor=tk.W, pady=(0, 5))
        
        self.text_input = scrolledtext.ScrolledText(
            input_frame, 
            wrap=tk.WORD, 
            font=('Helvetica', 11), 
            bg='#f8f9fa', 
            fg='#333333',
            borderwidth=1,
            relief=tk.SOLID
        )
        self.text_input.pack(fill=tk.BOTH, expand=True, pady=5)
        self.text_input.drop_target_register(DND_FILES)
        self.text_input.dnd_bind('<<Drop>>', self._handle_drop)
        
        btn_frame = ttk.Frame(input_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(btn_frame, text="Load Files", command=self._load_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Clear", command=self._clear_input).pack(side=tk.LEFT, padx=5)
        
        # Results Tab
        results_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(results_frame, text="Results")
        
        model_frame = ttk.LabelFrame(results_frame, text="AI Model Detection", padding=10)
        model_frame.pack(fill=tk.X, pady=5)
        
        self.model_vars = {}
        self.model_bars = {}
        for i, model in enumerate(['gpt4', 'gpt3', 'claude', 'deepseek']):
            frame = ttk.Frame(model_frame)
            frame.pack(fill=tk.X, pady=2)
            ttk.Label(frame, text=model.upper(), width=10, font=('Helvetica', 10, 'bold')).pack(side=tk.LEFT)
            self.model_vars[model] = tk.StringVar(value="0%")
            ttk.Label(frame, textvariable=self.model_vars[model], width=10).pack(side=tk.LEFT, padx=5)
            self.model_bars[model] = ttk.Progressbar(
                frame, 
                orient=tk.HORIZONTAL, 
                length=200, 
                mode='determinate',
                style='TProgressbar'
            )
            self.model_bars[model].pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        prob_frame = ttk.Frame(results_frame)
        prob_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(prob_frame, text="Overall AI Probability:", font=('Helvetica', 12, 'bold')).pack(side=tk.LEFT)
        self.prob_var = tk.StringVar(value="0%")
        ttk.Label(prob_frame, textvariable=self.prob_var, font=('Helvetica', 12, 'bold')).pack(side=tk.LEFT, padx=10)
        self.prob_bar = ttk.Progressbar(
            results_frame, 
            orient=tk.HORIZONTAL, 
            length=400, 
            mode='determinate',
            style='TProgressbar'
        )
        self.prob_bar.pack(fill=tk.X, pady=5)
        
        suggest_frame = ttk.LabelFrame(results_frame, text="Suggestions", padding=10)
        suggest_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.suggest_list = tk.Listbox(
            suggest_frame, 
            font=('Helvetica', 10), 
            bg='#f8f9fa', 
            fg='#333333',
            borderwidth=1,
            relief=tk.SOLID
        )
        scrollbar = ttk.Scrollbar(suggest_frame, orient=tk.VERTICAL, command=self.suggest_list.yview)
        self.suggest_list.configure(yscrollcommand=scrollbar.set)
        self.suggest_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Processing Tabs
        process_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(process_frame, text="Process")
        
        process_btn_frame = ttk.Frame(process_frame)
        process_btn_frame.pack(fill=tk.X)
        
        ttk.Button(process_btn_frame, text="Analyze", command=self._analyze_text).pack(side=tk.LEFT, padx=5)
        ttk.Button(process_btn_frame, text="Humanize", command=self._humanize_text).pack(side=tk.LEFT, padx=5)
        ttk.Button(process_btn_frame, text="Summarize", command=self._summarize_text).pack(side=tk.LEFT, padx=5)
        ttk.Button(process_btn_frame, text="Paraphrase", command=self._paraphrase_text).pack(side=tk.LEFT, padx=5)
        
        # Settings Tab
        settings_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(settings_frame, text="Settings")
        
        ttk.Label(settings_frame, text="Paraphrasing Diversity:", font=('Helvetica', 10, 'bold')).pack(anchor=tk.W)
        self.diversity_var = tk.DoubleVar(value=0.7)
        diversity_scale = ttk.Scale(
            settings_frame, 
            from_=0.1, 
            to=0.9, 
            orient=tk.HORIZONTAL, 
            variable=self.diversity_var,
            style='TScale'
        )
        diversity_scale.pack(fill=tk.X, pady=5)
        ttk.Label(settings_frame, textvariable=self.diversity_var, font=('Helvetica', 10)).pack(anchor=tk.W)
        
        ttk.Label(settings_frame, text="Humanization Level:", font=('Helvetica', 10, 'bold')).pack(anchor=tk.W, pady=5)
        self.humanization_var = tk.DoubleVar(value=0.5)
        humanization_scale = ttk.Scale(
            settings_frame, 
            from_=0.1, 
            to=1.0, 
            orient=tk.HORIZONTAL, 
            variable=self.humanization_var,
            style='TScale'
        )
        humanization_scale.pack(fill=tk.X, pady=5)
        ttk.Label(settings_frame, textvariable=self.humanization_var, font=('Helvetica', 10)).pack(anchor=tk.W)
        
        # Learning Tab
        learn_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(learn_frame, text="Learn")
        
        ttk.Button(learn_frame, text="Upload Training Documents", command=self._load_learning_files).pack(anchor=tk.W)
        self.learn_status = tk.StringVar(value="No documents trained")
        ttk.Label(learn_frame, textvariable=self.learn_status, font=('Helvetica', 10)).pack(anchor=tk.W, pady=5)
        
        # History Tab
        history_frame = ttk.Frame(self.notebook, padding=10)
        self.notebook.add(history_frame, text="History")
        
        self.history_list = tk.Listbox(
            history_frame, 
            font=('Helvetica', 10), 
            bg='#f8f9fa', 
            fg='#333333',
            borderwidth=1,
            relief=tk.SOLID
        )
        scrollbar = ttk.Scrollbar(history_frame, orient=tk.VERTICAL, command=self.history_list.yview)
        self.history_list.configure(yscrollcommand=scrollbar.set)
        self.history_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.history_list.bind('<<ListboxSelect>>', self._view_history)
        
        # Status and Progress
        status_frame = ttk.Frame(self.root)
        status_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=5)
        
        self.status_var = tk.StringVar(value="Ready")
        status_label = ttk.Label(
            status_frame, 
            textvariable=self.status_var, 
            relief=tk.SUNKEN, 
            background='#e9ecef', 
            foreground='#333333',
            padding=5
        )
        status_label.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            status_frame, 
            variable=self.progress_var, 
            maximum=100,
            style='TProgressbar'
        )
        self.progress_bar.pack(side=tk.BOTTOM, fill=tk.X, pady=5)
        
        # Real-time log display
        log_frame = ttk.LabelFrame(self.root, text="Activity Log", padding=10)
        log_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=5)
        
        self.log_text = scrolledtext.ScrolledText(
            log_frame,
            wrap=tk.WORD,
            height=5,
            font=('Helvetica', 10),
            bg='#f8f9fa',
            fg='#333333',
            borderwidth=1,
            relief=tk.SOLID
        )
        self.log_text.pack(fill=tk.X)
        self.log_text.configure(state='disabled')
        
        # Configure scale style
        style.configure('TScale', background='#ffffff', troughcolor='#dee2e6')
    
    def _log_message(self, message):
        """Log messages to the GUI log display"""
        self.log_text.configure(state='normal')
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.log_text.configure(state='disabled')
    
    def _animate_progress(self):
        """Animate progress bar during processing"""
        if self.progress_var.get() < 100:
            self.progress_var.set(self.progress_var.get() + 2)
            self.progress_animation_id = self.root.after(50, self._animate_progress)
        else:
            self.progress_var.set(100)
    
    def _stop_animation(self):
        """Stop progress bar animation"""
        if self.progress_animation_id:
            self.root.after_cancel(self.progress_animation_id)
            self.progress_animation_id = None
    
    def _handle_drop(self, event):
        """Handle drag-and-drop files"""
        files = self.root.splitlist(event.data)
        self.file_queue.extend(files)
        self._log_message(f"Dropped {len(files)} file(s)")
        self._process_file_queue()
    
    def _load_files(self):
        """Load multiple files (.txt, .docx, .pdf)"""
        file_paths = filedialog.askopenfilenames(filetypes=[
            ("Text files", "*.txt"),
            ("Word documents", "*.docx"),
            ("PDF files", "*.pdf"),
            ("All files", "*.*")
        ])
        if file_paths:
            self.file_queue.extend(file_paths)
            self._log_message(f"Selected {len(file_paths)} file(s) for loading")
            self._process_file_queue()
    
    def _load_learning_files(self):
        """Load human-written documents for training"""
        file_paths = filedialog.askopenfilenames(filetypes=[
            ("Word documents", "*.docx"),
            ("PDF files", "*.pdf"),
            ("All files", "*.*")
        ])
        if file_paths:
            self.status_var.set("Processing learning documents...")
            self._log_message("Started processing learning documents")
            self.progress_var.set(0)
            self._animate_progress()
            
            def load_task():
                texts = []
                total = len(file_paths)
                for i, file_path in enumerate(file_paths):
                    try:
                        if file_path.endswith('.docx'):
                            doc = docx.Document(file_path)
                            texts.append('\n'.join([para.text for para in doc.paragraphs]))
                            self._log_message(f"Loaded {os.path.basename(file_path)}")
                        elif file_path.endswith('.pdf'):
                            with pdfplumber.open(file_path) as pdf:
                                texts.append('\n'.join([page.extract_text() or '' for page in pdf.pages]))
                            self._log_message(f"Loaded {os.path.basename(file_path)}")
                        self.progress_var.set((i + 1) / total * 100)
                    except Exception as e:
                        self._log_message(f"Error loading {file_path}: {str(e)}")
                return texts
            
            def callback(texts):
                if texts:
                    self.detector.learn_from_documents(texts)
                    self.learn_status.set(f"Trained on {len(texts)} documents")
                    self._log_message(f"Training complete on {len(texts)} documents")
                else:
                    self.status_var.set("No valid documents loaded")
                    self._log_message("No valid documents were loaded")
                self._stop_animation()
                self.progress_var.set(0)
            
            self.task_manager.add_task(load_task, callback)
    
    def _process_file_queue(self):
        """Process files in the queue"""
        if not self.file_queue:
            self.status_var.set("Ready")
            self._log_message("File queue processing complete")
            self._stop_animation()
            return
        
        file_path = self.file_queue.pop(0)
        self.status_var.set(f"Loading {os.path.basename(file_path)}...")
        self._log_message(f"Loading file: {os.path.basename(file_path)}")
        
        def load_task():
            try:
                if file_path.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return f.read()
                elif file_path.endswith('.docx'):
                    doc = docx.Document(file_path)
                    return '\n'.join([para.text for para in doc.paragraphs])
                elif file_path.endswith('.pdf'):
                    with pdfplumber.open(file_path) as pdf:
                        return '\n'.join([page.extract_text() or '' for page in pdf.pages])
                else:
                    return "Error: Unsupported file format"
            except Exception as e:
                return f"Error: {str(e)}"
        
        def callback(content):
            if content.startswith("Error:"):
                messagebox.showerror("Error", content[6:])
                self._log_message(f"Error loading file: {content[6:]}")
            else:
                self.text_input.delete(1.0, tk.END)
                self.text_input.insert(1.0, content)
                self.history.append(('Loaded', content[:100] + '...' if len(content) > 100 else content))
                self._update_history()
                self._log_message(f"Successfully loaded {os.path.basename(file_path)}")
            self._process_file_queue()
        
        self.task_manager.add_task(load_task, callback)
    
    def _clear_input(self):
        """Clear input and reset results"""
        self.text_input.delete(1.0, tk.END)
        self.file_queue = []
        for model in self.model_vars:
            self.model_vars[model].set("0%")
            self.model_bars[model]['value'] = 0
        self.prob_var.set("0%")
        self.prob_bar['value'] = 0
        self.suggest_list.delete(0, tk.END)
        self.status_var.set("Input cleared")
        self._log_message("Input cleared and results reset")
        self._stop_animation()
    
    def _analyze_text(self):
        """Analyze text for AI patterns"""
        text = self.text_input.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Warning", "Please enter text to analyze")
            self._log_message("Warning: No text provided for analysis")
            return
        
        self.status_var.set("Analyzing content...")
        self._log_message("Starting text analysis")
        self.progress_var.set(0)
        self._animate_progress()
        
        def analysis_task():
            return self.detector.analyze_text(text)
        
        def callback(results):
            self._stop_animation()
            if 'error' in results:
                messagebox.showerror("Error", results['error'])
                self.status_var.set("Error in analysis")
                self._log_message(f"Analysis error: {results['error']}")
                return
            
            for model, score in results['model_scores'].items():
                self.model_vars[model].set(f"{score:.1f}%")
                self.model_bars[model]['value'] = score
            self.prob_var.set(f"{results['ai_probability']:.1f}%")
            self.prob_bar['value'] = results['ai_probability']
            self.suggest_list.delete(0, tk.END)
            for suggestion in results['suggestions']:
                self.suggest_list.insert(tk.END, suggestion)
            self.history.append(('Analyzed', f"AI Probability: {results['ai_probability']:.1f}%"))
            self._update_history()
            self.status_var.set(f"Analysis complete - Likely {results['likely_model'] or 'unknown'}")
            self._log_message(f"Analysis complete: {results['ai_probability']:.1f}% AI probability")
            self.progress_var.set(0)
        
        self.task_manager.add_task(analysis_task, callback)
    
    def _humanize_text(self):
        """Humanize the text in background"""
        text = self.text_input.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Warning", "Please enter text to humanize")
            self._log_message("Warning: No text provided for humanization")
            return
        
        self.status_var.set("Humanizing text...")
        self._log_message("Starting text humanization")
        self.progress_var.set(0)
        self._animate_progress()
        
        def humanize_task():
            return self.detector.humanize_text(text, self.diversity_var.get())
        
        def callback(result):
            self._stop_animation()
            if result.startswith("Error:"):
                messagebox.showerror("Error", result[6:])
                self.status_var.set("Error in humanization")
                self._log_message(f"Humanization error: {result[6:]}")
                return
            
            win = tk.Toplevel(self.root)
            win.title("Humanized Text")
            win.geometry("800x600")
            win.configure(bg='#ffffff')
            
            frame = ttk.Frame(win, padding=10)
            frame.pack(fill=tk.BOTH, expand=True)
            
            ttk.Label(frame, text="Humanized Version:", font=('Helvetica', 12, 'bold')).pack(anchor=tk.W)
            
            text_box = scrolledtext.ScrolledText(
                frame, 
                wrap=tk.WORD, 
                font=('Helvetica', 11), 
                bg='#f8f9fa', 
                fg='#333333',
                borderwidth=1,
                relief=tk.SOLID
            )
            text_box.pack(fill=tk.BOTH, expand=True, pady=5)
            text_box.insert(1.0, result)
            
            btn_frame = ttk.Frame(frame)
            btn_frame.pack(fill=tk.X, pady=5)
            
            ttk.Button(btn_frame, text="Copy", command=lambda: self._copy_to_clipboard(result)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Replace", command=lambda: self._replace_text(result, win)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Close", command=win.destroy).pack(side=tk.RIGHT)
            
            self.history.append(('Humanized', result[:100] + '...' if len(result) > 100 else result))
            self._update_history()
            self.status_var.set("Humanization complete")
            self._log_message("Text humanization completed")
            self.progress_var.set(0)
        
        self.task_manager.add_task(humanize_task, callback)
    
    def _summarize_text(self):
        """Summarize the input text"""
        text = self.text_input.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Warning", "Please enter text to summarize")
            self._log_message("Warning: No text provided for summarization")
            return
        
        self.status_var.set("Summarizing text...")
        self._log_message("Starting text summarization")
        self.progress_var.set(0)
        self._animate_progress()
        
        def summarize_task():
            return self.detector.summarize_text(text)
        
        def callback(result):
            self._stop_animation()
            if result.startswith("Error:"):
                messagebox.showerror("Error", result[6:])
                self.status_var.set("Error in summarization")
                self._log_message(f"Summarization error: {result[6:]}")
                return
            
            win = tk.Toplevel(self.root)
            win.title("Summarized Text")
            win.geometry("800x600")
            win.configure(bg='#ffffff')
            
            frame = ttk.Frame(win, padding=10)
            frame.pack(fill=tk.BOTH, expand=True)
            
            ttk.Label(frame, text="Summary:", font=('Helvetica', 12, 'bold')).pack(anchor=tk.W)
            
            text_box = scrolledtext.ScrolledText(
                frame, 
                wrap=tk.WORD, 
                font=('Helvetica', 11), 
                bg='#f8f9fa', 
                fg='#333333',
                borderwidth=1,
                relief=tk.SOLID
            )
            text_box.pack(fill=tk.BOTH, expand=True, pady=5)
            text_box.insert(1.0, result)
            
            btn_frame = ttk.Frame(frame)
            btn_frame.pack(fill=tk.X, pady=5)
            
            ttk.Button(btn_frame, text="Copy", command=lambda: self._copy_to_clipboard(result)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Replace", command=lambda: self._replace_text(result, win)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Close", command=win.destroy).pack(side=tk.RIGHT)
            
            self.history.append(('Summarized', result[:100] + '...' if len(result) > 100 else result))
            self._update_history()
            self.status_var.set("Summarization complete")
            self._log_message("Text summarization completed")
            self.progress_var.set(0)
        
        self.task_manager.add_task(summarize_task, callback)
    
    def _paraphrase_text(self):
        """Paraphrase the input text"""
        text = self.text_input.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Warning", "Please enter text to paraphrase")
            self._log_message("Warning: No text provided for paraphrasing")
            return
        
        self.status_var.set("Paraphrasing text...")
        self._log_message("Starting text paraphrasing")
        self.progress_var.set(0)
        self._animate_progress()
        
        def paraphrase_task():
            return self.detector.paraphrase_text(text, self.diversity_var.get())
        
        def callback(result):
            self._stop_animation()
            if result.startswith("Error:"):
                messagebox.showerror("Error", result[6:])
                self.status_var.set("Error in paraphrasing")
                self._log_message(f"Paraphrasing error: {result[6:]}")
                return
            
            win = tk.Toplevel(self.root)
            win.title("Paraphrased Text")
            win.geometry("800x600")
            win.configure(bg='#ffffff')
            
            frame = ttk.Frame(win, padding=10)
            frame.pack(fill=tk.BOTH, expand=True)
            
            ttk.Label(frame, text="Paraphrased Version:", font=('Helvetica', 12, 'bold')).pack(anchor=tk.W)
            
            text_box = scrolledtext.ScrolledText(
                frame, 
                wrap=tk.WORD, 
                font=('Helvetica', 11), 
                bg='#f8f9fa', 
                fg='#333333',
                borderwidth=1,
                relief=tk.SOLID
            )
            text_box.pack(fill=tk.BOTH, expand=True, pady=5)
            text_box.insert(1.0, result)
            
            btn_frame = ttk.Frame(frame)
            btn_frame.pack(fill=tk.X, pady=5)
            
            ttk.Button(btn_frame, text="Copy", command=lambda: self._copy_to_clipboard(result)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Replace", command=lambda: self._replace_text(result, win)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Close", command=win.destroy).pack(side=tk.RIGHT)
            
            self.history.append(('Paraphrased', result[:100] + '...' if len(result) > 100 else result))
            self._update_history()
            self.status_var.set("Paraphrasing complete")
            self._log_message("Text paraphrasing completed")
            self.progress_var.set(0)
        
        self.task_manager.add_task(paraphrase_task, callback)
    
    def _copy_to_clipboard(self, text):
        """Copy text to clipboard"""
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.status_var.set("Copied to clipboard")
        self._log_message("Text copied to clipboard")
    
    def _replace_text(self, text, window):
        """Replace original text with processed version"""
        self.text_input.delete(1.0, tk.END)
        self.text_input.insert(1.0, text)
        window.destroy()
        self.status_var.set("Text replaced")
        self._log_message("Text replaced in input area")
    
    def _update_history(self):
        """Update the history list"""
        self.history_list.delete(0, tk.END)
        for i, (action, content) in enumerate(self.history[-10:]):  # Show last 10 entries
            self.history_list.insert(tk.END, f"{action}: {content}")
        self._log_message("History updated")
    
    def _view_history(self, event):
        """View selected history entry"""
        selection = self.history_list.curselection()
        if selection:
            index = selection[0]
            action, content = self.history[index]
            win = tk.Toplevel(self.root)
            win.title(f"{action} Content")
            win.geometry("800x600")
            win.configure(bg='#ffffff')
            
            frame = ttk.Frame(win, padding=10)
            frame.pack(fill=tk.BOTH, expand=True)
            
            ttk.Label(frame, text=f"{action} Content:", font=('Helvetica', 12, 'bold')).pack(anchor=tk.W)
            
            text_box = scrolledtext.ScrolledText(
                frame, 
                wrap=tk.WORD, 
                font=('Helvetica', 11), 
                bg='#f8f9fa', 
                fg='#333333',
                borderwidth=1,
                relief=tk.SOLID
            )
            text_box.pack(fill=tk.BOTH, expand=True, pady=5)
            text_box.insert(1.0, content)
            
            ttk.Button(frame, text="Close", command=win.destroy).pack(side=tk.RIGHT, pady=5)
            self._log_message(f"Viewing {action} history entry")
    
    def on_close(self):
        """Cleanup on window close"""
        self._stop_animation()
        self.task_manager.shutdown()
        self._log_message("Application closed")
        self.root.destroy()

def main():
    app = AIDetectorGUI()
    app.root.protocol("WM_DELETE_WINDOW", app.on_close)
    app.root.mainloop()

if __name__ == "__main__":
    main()